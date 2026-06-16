"""Export utilities for generating PDF and Word documents from OCR text"""
import io
import re
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BRAND_NAME = "HospAI"
BRAND_TAGLINE = "AI-Driven Healthcare Optimization"


def parse_markdown_line(line):
    """Parse a line and return (type, content, is_bold_segments)
    Types: header, bullet, table_row, text
    """
    line = line.strip()
    if not line:
        return ("empty", "", [])

    header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
    if header_match:
        return ("header", header_match.group(2).replace("**", ""), [])

    bullet_match = re.match(r"^[\•\-\*]\s+(.+)$", line)
    if bullet_match:
        content = bullet_match.group(1)
        return ("bullet", content, parse_bold_segments(content))

    if line.startswith("|") and line.endswith("|"):
        if re.match(r"^\|[\s\-\|:]+\|$", line):
            return ("table_sep", "", [])
        cells = [c.strip().replace("**", "") for c in line.strip("|").split("|")]
        return ("table_row", cells, [])

    return ("text", line, parse_bold_segments(line))


def parse_bold_segments(text):
    """Parse text and return list of (text, is_bold) tuples"""
    segments = []
    pattern = r"\*\*([^*]+)\*\*"
    last_end = 0

    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            segments.append((text[last_end:match.start()], False))
        segments.append((match.group(1), True))
        last_end = match.end()

    if last_end < len(text):
        segments.append((text[last_end:], False))

    if not segments:
        segments = [(text, False)]

    return segments


def generate_pdf(patient_name: str, doc_type: str, ocr_text: str, date_str: str = None) -> bytes:
    """Generate PDF with OCR extracted text - properly formatted"""
    if date_str is None:
        date_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    pdf = FPDF()
    left_margin = 20
    right_margin = 20
    pdf.set_left_margin(left_margin)
    pdf.set_right_margin(right_margin)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    page_width = pdf.w
    effective_width = page_width - left_margin - right_margin

    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(20, 61, 94)
    pdf.cell(effective_width, 12, BRAND_NAME, ln=True, align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(effective_width, 6, BRAND_TAGLINE, ln=True, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(effective_width, 10, "Medical Document Report", ln=True, align="C")
    pdf.ln(6)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(45, 8, "Patient Name:")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, patient_name, ln=True)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(45, 8, "Document Type:")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, doc_type.replace("_", " ").title(), ln=True)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(45, 8, "Date:")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, date_str, ln=True)

    pdf.ln(8)
    pdf.set_draw_color(80, 80, 80)
    pdf.set_line_width(0.5)
    pdf.line(left_margin, pdf.get_y(), page_width - right_margin, pdf.get_y())
    pdf.ln(10)

    lines = ocr_text.split("\n")

    for line in lines:
        line_type, content, segments = parse_markdown_line(line)

        if line_type == "empty":
            pdf.ln(4)
            continue

        if line_type == "header":
            pdf.ln(6)
            pdf.set_draw_color(150, 150, 150)
            pdf.set_line_width(0.3)
            pdf.line(left_margin, pdf.get_y(), page_width - right_margin, pdf.get_y())
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 13)
            safe_text = content.encode("latin-1", errors="replace").decode("latin-1")
            pdf.cell(effective_width, 8, safe_text, ln=True)
            pdf.ln(3)
            continue

        if line_type == "bullet":
            pdf.set_x(left_margin + 8)
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(5, 6, chr(149))
            for text, is_bold in segments:
                text = text.replace("**", "")
                safe_text = text.encode("latin-1", errors="replace").decode("latin-1")
                pdf.set_font("Helvetica", "B" if is_bold else "", 10)
                pdf.write(6, safe_text)
            pdf.ln(6)
            continue

        if line_type == "table_row":
            pdf.set_font("Helvetica", "", 10)
            col_width = effective_width / max(len(content), 1)
            for cell in content:
                safe_text = cell.encode("latin-1", errors="replace").decode("latin-1")
                pdf.cell(col_width, 7, safe_text[:30], border=0)
            pdf.ln(7)
            continue

        if line_type == "table_sep":
            continue

        if line_type == "text":
            for text, is_bold in segments:
                text = text.replace("**", "")
                safe_text = text.encode("latin-1", errors="replace").decode("latin-1")
                pdf.set_font("Helvetica", "B" if is_bold else "", 10)
                pdf.write(6, safe_text)
            pdf.ln(6)

    return bytes(pdf.output())


def generate_word(patient_name: str, doc_type: str, ocr_text: str, date_str: str = None) -> bytes:
    """Generate Word document with OCR extracted text - properly formatted"""
    if date_str is None:
        date_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    doc = Document()

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand.add_run(BRAND_NAME)
    brand_run.bold = True
    brand_run.font.size = Pt(24)
    brand_run.font.color.rgb = RGBColor(20, 61, 94)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(BRAND_TAGLINE)
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    title = doc.add_heading("Medical Document Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"

    cells = table.rows[0].cells
    cells[0].text = "Patient Name"
    cells[1].text = patient_name

    cells = table.rows[1].cells
    cells[0].text = "Document Type"
    cells[1].text = doc_type.replace("_", " ").title()

    cells = table.rows[2].cells
    cells[0].text = "Date"
    cells[1].text = date_str

    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    lines = ocr_text.split("\n")

    for line in lines:
        line_type, content, segments = parse_markdown_line(line)

        if line_type == "empty":
            doc.add_paragraph()
            continue

        if line_type == "header":
            doc.add_heading(content, level=2)
            continue

        if line_type == "bullet":
            para = doc.add_paragraph(style="List Bullet")
            for text, is_bold in segments:
                text = text.replace("**", "")
                run = para.add_run(text)
                run.bold = is_bold
            continue

        if line_type == "table_row":
            para = doc.add_paragraph()
            para.add_run("    ".join(content))
            continue

        if line_type == "table_sep":
            continue

        if line_type == "text":
            para = doc.add_paragraph()
            for text, is_bold in segments:
                text = text.replace("**", "")
                run = para.add_run(text)
                run.bold = is_bold

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
